from helpers.logger_functions import info_log
from helpers.helper_functions import clean_and_parse_json
import  json
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from linkedin_bot import OPENAI_API_KEY

def get_chatgpt_response(resume_text=None, job_description='example job', job_title='example title', job_company='example company'):
    client = OpenAI(api_key=OPENAI_API_KEY)
    if not resume_text:
      with open('resume.txt', 'r') as f:
        resume_text = f.read()
      with open('sample_description.txt', 'r') as f:
        job_description = f.read()
    prompt = 'Using the following text between the brackets as my resume {'+resume_text+'}, rewrite my resume using the following job description between the following brackets {'+job_description+'}.     Use Markdown formatting for sections (##), bold job titles, and bullet points for skills and experience.     Return ONLY valid JSON with these keys:    - name (string)    - contact_info (string)    - summary (string, plain text)    - skills (object with keys: languages, frameworks, frontend, backend, devops, database, testing — each is a list of strings)    - experience (list of objects with company, role, dates, description as list of bullet points)    - education (list of objects with institution, degree, dates)    Make descriptions concise and achievements-focused. ONLY output the JSON object, no explanations.'

    info_log('Accessing OpenAI API...')
    response = client.chat.completions.create(
      model="gpt-4o-mini",
      messages=[{"role": "user", "content": prompt}],
      temperature=0,
  )
    
    json_text = response.choices[0].message.content

    try:
        resume_data = clean_and_parse_json(json_text)
    except json.JSONDecodeError as e:
        print("Failed to parse JSON:", e)
        print("GPT response was:", json_text)
        resume_data = None

    doc = Document()

    def add_heading(text, level=1):
        para = doc.add_heading(text, level=level)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_paragraph(text, bold=False, italic=False, font_size=11, indent=False):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        if indent:
            para.paragraph_format.left_indent = Inches(0.25)
        return para

    def add_bullet_point(text):
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(text)
        run.font.size = Pt(11)
        return para

    # Header: Name centered, blue, large font
    name_para = doc.add_paragraph()
    run = name_para.add_run(resume_data["name"])
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(46, 134, 193)  # Blue
    name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Contact info centered, smaller font
    contact_para = doc.add_paragraph(resume_data["contact_info"])
    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_para.style.font.size = Pt(10)

    doc.add_paragraph()  # blank line

    # Summary
    add_heading("Summary", level=2)
    add_paragraph(resume_data["summary"], font_size=11)

    # Skills (inline comma-separated)
    add_heading("Skills", level=2)
    for category, items in resume_data["skills"].items():
      # Capitalize category title
      add_paragraph(category.capitalize() + ":", bold=True, font_size=11)
      # Join items with commas
      add_paragraph(", ".join(items), font_size=11, indent=True)

    # Experience
    add_heading("Experience", level=2)
    for exp in resume_data["experience"]:
        add_paragraph(f"{exp['role']} — {exp['company']} ({exp['dates']})", bold=True, font_size=12)
        for bullet in exp["description"]:
            add_bullet_point(bullet)

    # Education
    add_heading("Education", level=2)
    for edu in resume_data["education"]:
        add_paragraph(f"{edu['degree']}", bold=True, font_size=12)
        add_paragraph(f"{edu['institution']} ({edu['dates']})", font_size=11)
        doc.add_paragraph()  # spacing between education entries

    info_log('Got markdown text')
    info_log("Response extracted.")
    # Build new resume
    doc_name = f'Ryan Mooney Resume - {job_title} at {job_company}.docx'
    current_dir = os.getcwd()
    filepath = os.path.join(current_dir, 'resumes', doc_name)
    doc.save(filepath)
    info_log(f'Resume saved as {doc_name}.docx.')
    return filepath